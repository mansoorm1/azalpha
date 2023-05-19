'''
AZALPHA
Some functions for azarbaijani alphabet translitration

'''

from docx import Document
import os
import pandas as pd
import operator

word_boder = ' .;,?!'  

def main():

	print ('AZALPHA is a set of functions, supposed to be used higher level applications')
	print ('Consult main help for more information')

	
	return 

# Python3 code to demonstrate working of
# Sort Dictionary by Key Lengths
# Using len() + sort() + dictionary comprehension + items()
 
def get_len(key):
    return len(key[0])
 
# sorting using sort()
# external to render logic
def sort_dict(dict,desc=True):
	dict_list = list(dict.items())
	dict_list.sort(key = get_len, reverse=desc)
	# reordering to dictionary
	dict_sorted = {ele[0] : ele[1]  for ele in dict_list}
	return dict_sorted
 

def print_terms(terms):
	for term in terms:
		try:
			print( term.lower().encode('utf-8'), terms[term].encode('utf-8'))
		except:
			print ('Some error')
			
	
#Wrtie the dictionary to file:
def write_to_file(terms, fname):
	with open(fname, "w", encoding="utf-8") as f:
		for k, v in terms.items():
			f.write(k.strip()+ "\t " + v.strip() + "\n")

def map_term(df , word , col1 , col2):
	match = '' 
	n = len(word)
	if n == 0:
		match = ''
	elif n == 1:
		rows = df.loc[df[col1] == word]
		if len(rows) > 0:
			match   = rows[col2].values[0]
		else:
			match = word #unknown character is there
			try:
				print('Character not found. default used' , word)
			except:
				print('Character not found. default used' , word.encode('utf-8'))
			
	else:
		found = False
		for m in reversed(range(n+1)) :
			if found: break
			for k in range(0, n-m+1):
				if found: break
				subword = word[k:k+m]
				rows = df.loc[df[col1] == subword]
				if len(rows) > 0:
					match_mid   = rows[col2].values[0]
					found = True
				elif m == 1:
					match_mid = subword #a single alphabet character is not found! take the original 
					found = True
				if found:
					left  = word[0:k]
					right = word[k+m:]
					match_left  = map_term(df, left,  col1, col2)
					match_right = map_term(df, right, col1, col2)
					match = match_left + match_mid + match_right
					found = True
					
	return match 


'''
Get all the terms in document as a dictionary.
'''
def get_doc_terms(doc):		
	#create an empty dictionary
	terms = {}
	
	#scan all paragraphs and tables and collect terms
	for paragraph in doc.paragraphs:
		for run in paragraph.runs:
			items = run.text.split()
			for item in items:
				item = item.strip().lower()
				if len(item) > 0:
					terms[item] =''

	for table in doc.tables:
		for col in table.columns:
			for cell in col.cells:
				for paragraph in cell.paragraphs:
					for run in paragraph.runs:
						items = run.text.split()
						for item in items:
							item = item.strip().lower()
							if len(item) > 0:
								terms[item] =''
	
	#return collected terms
	terms = dict(sorted(terms.items()))
	terms = sort_dict(terms)
	return terms



#Map all the terms to the target alphabet
def map_terms(terms, df, col1 , col2):
	for key, value in terms.items():
		value      = map_term(df, key, col1, col2)
		terms[key] = value.strip().lower()
	
	return terms

#Map a single term:
def lookup_term(terms, term):
	item2 = ''
	item  = term.strip().lower()
	if len(item) > 0 :
		item2 = terms.get(item, "#ERROR")
	
	return item2


#Map the document
def map_doc(doc, terms):
	#scan all paragraphs and tables and collect terms
	for key , value in terms.items():
		print ('key = ' , key , 'value = ' , value)
		for paragraph in doc.paragraphs:
			for run in paragraph.runs:
				run.text = run.text.lower() 
				run.text = run.text.replace(key , value)

		for table in doc.tables:
			for col in table.columns:
				for cell in col.cells:
					for paragraph in cell.paragraphs:
						for run in paragraph.runs:
							run.text = run.text.lower() 
							run.text = run.text.replace(key , value)
							

	return doc

#Map the document
def map_doc2(doc, terms):
	#scan all paragraphs and tables and collect terms
	for paragraph in doc.paragraphs:
		for run in paragraph.runs:
			newtext = ''
			items = run.text.split()
			for item in items:
				item = item.strip().lower()
				item = lookup_term(terms, item)
				newtext = newtext + ' ' +item
			run.text = newtext

	for table in doc.tables:
		for col in table.columns:
			for cell in col.cells:
				for paragraph in cell.paragraphs:
					for run in paragraph.runs:
						newtext = ''
						items = run.text.split()
						for item in items:
							item = item.strip().lower()
							item = lookup_term(terms, item)
							newtext = newtext + ' ' +item
						run.text = newtext

	return doc


def default_log_message(msg_to_log):
	print(msg_to_log)	
	

def arla(srcdoc, dstdoc, sozluk, direction, log_msg=default_log_message):
	if direction == "ar2la":
		col1 = "Arabic"
		col2 = "Latin"
	elif direction == "la2ar": 
		col2 = "Arabic"
		col1 = "Latin"
	else:
		log_msg ('Unknown direction{}'.format(direction))
		return
	
	
	
	df = pd.read_excel(sozluk)
	df = df.reset_index()  # make sure indexes pair with number of rows
	
	document = Document(srcdoc)

	
	#collect and map the terms:
	log_msg('Collecting document terms ...' )
	terms = get_doc_terms(document)
	write_to_file(terms , './terms.csv')
	log_msg('{} words found. see the <terms.csv> file'.format(len(terms)) )
	
	log_msg('Transliterate document terms ...' )
	terms = map_terms(terms, df , col1,  col2, )
	write_to_file(terms , './terms2.csv')

	log_msg('Transliterate document ...' )
	document = map_doc(document, terms)
	log_msg('Saving document ...' )
	document.save(dstdoc)
	log_msg('Finished!' )

	
	return 


if __name__ == '__main__':
	main()
